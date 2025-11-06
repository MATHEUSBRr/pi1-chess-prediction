# generate_pi1_chess_pgn.py
# Lê um arquivo PGN local (partidas_matheus.pgn), extrai features, treina RandomForest + XGBoost/GradientBoosting,
# salva figuras, modelos e gera um relatório .docx pronto para entregar.
#
# Requer: python-chess, pandas, numpy, scikit-learn, matplotlib, python-docx, joblib
# pip install python-chess pandas numpy scikit-learn matplotlib python-docx joblib xgboost

import os, io, math, time
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import joblib
import chess.pgn

from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.metrics import classification_report, confusion_matrix, accuracy_score
from sklearn.dummy import DummyClassifier

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ---------- Config ----------
PGN_FILE = "partidas_matheus.pgn"
OUT_DIR = "pi1_pgn_output"
FIG_DIR = os.path.join(OUT_DIR, "figs")
MODEL_DIR = os.path.join(OUT_DIR, "models")
os.makedirs(FIG_DIR, exist_ok=True)
os.makedirs(MODEL_DIR, exist_ok=True)

# ---------- Helpers ----------
def load_pgn_blocks_from_file(path):
    """Lê um arquivo PGN que pode conter várias partidas e retorna lista de PGN strings."""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    blocks = []
    if "\n\n\n" in text:
        parts = [b.strip() for b in text.split("\n\n\n") if b.strip()]
    else:
        parts = []
        current = []
        for ln in text.splitlines():
            if ln.strip().startswith("[Event") and current:
                parts.append("\n".join(current).strip())
                current = [ln]
            else:
                current.append(ln)
        if current:
            parts.append("\n".join(current).strip())
    for p in parts:
        if p.startswith("[") and len(p) > 80:
            blocks.append(p)
    return blocks

def extract_features_from_pgn(pgn_text):
    """Extrai features relevantes de um PGN (usa python-chess)."""
    try:
        game = chess.pgn.read_game(io.StringIO(pgn_text))
        if game is None:
            return None
    except Exception:
        return None

    tags = game.headers
    # ratings
    try:
        wr = int(tags.get("WhiteElo", 0))
    except:
        wr = 0
    try:
        br = int(tags.get("BlackElo", 0))
    except:
        br = 0
    res_tag = tags.get("Result", "*")
    if res_tag == "1-0":
        res = "white"
    elif res_tag == "0-1":
        res = "black"
    else:
        res = "draw"
    opening = tags.get("Opening", tags.get("ECO", "Unknown"))
    tc = tags.get("TimeControl", "")

    node = game
    half_moves = 0
    while node.variations:
        node = node.variations[0]
        half_moves += 1
    full_moves = math.ceil(half_moves / 2)

    board = game.end().board()
    vals = {'P':1,'N':3,'B':3,'R':5,'Q':9}
    white_mat = 0
    black_mat = 0
    for _, piece in board.piece_map().items():
        v = vals.get(piece.symbol().upper(), 0)
        if piece.color:
            white_mat += v
        else:
            black_mat += v
    material_diff = white_mat - black_mat

    return {
        "white_rating": wr,
        "black_rating": br,
        "rating_diff": wr - br,
        "moves": full_moves,
        "opening": opening,
        "time_control": tc,
        "material_diff": material_diff,
        "result": res,
        "pgn": pgn_text
    }

# ---------- Load PGN file ----------
if not os.path.exists(PGN_FILE):
    print(f"Arquivo PGN não encontrado: {PGN_FILE}. Coloque o arquivo na mesma pasta deste script.")
    raise SystemExit(1)

print("Lendo PGN local:", PGN_FILE)
pgn_blocks = load_pgn_blocks_from_file(PGN_FILE)
print(f"Blocos PGN detectados: {len(pgn_blocks)}")

all_feats = []
for block in pgn_blocks:
    feats = extract_features_from_pgn(block)
    if feats:
        all_feats.append(feats)

if not all_feats:
    print("Nenhuma partida válida encontrada no PGN. Verifique o arquivo.")
    raise SystemExit(1)

df = pd.DataFrame(all_feats)

df['abs_rating_diff'] = df['rating_diff'].abs()
df['moves_log'] = np.log1p(df['moves'])
df['white_higher'] = (df['rating_diff'] > 0).astype(int)

csv_path = os.path.join(OUT_DIR, "chess_pgn_data.csv")
df.to_csv(csv_path, index=False)
print("CSV salvo em:", csv_path)

vals = df['result'].value_counts().reindex(['white','draw','black']).fillna(0)
plt.figure(figsize=(6,4)); plt.bar(vals.index, vals.values); plt.title('Distribuição de resultados (PGN)'); plt.xlabel('Resultado'); plt.ylabel('Contagem'); plt.tight_layout()
fig1 = os.path.join(FIG_DIR, "dist_results_pgn.png"); plt.savefig(fig1); plt.close()

plt.figure(figsize=(8,4))
groups = [df[df['result']==c]['rating_diff'] for c in ['white','draw','black']]
plt.boxplot(groups, labels=['white','draw','black'])
plt.title('Rating diff por resultado (PGN)'); plt.tight_layout()
fig2 = os.path.join(FIG_DIR, "boxplot_ratingdiff_pgn.png"); plt.savefig(fig2); plt.close()

plt.figure(figsize=(6,4)); plt.hist(df['moves'], bins=15); plt.title('Número de lances (PGN)'); plt.tight_layout()
fig3 = os.path.join(FIG_DIR, "hist_moves_pgn.png"); plt.savefig(fig3); plt.close()

# ---------- Prepare data for modeling ----------
features = ['white_rating','black_rating','rating_diff','abs_rating_diff','moves','moves_log','time_control','opening','material_diff','white_higher']
X = df[features].copy()
y = df['result'].copy()

if len(y.unique()) < 2:
    print("Atenção: menos de 2 classes encontradas nos dados. Os modelos podem não treinar corretamente.")

strat = y if len(y.unique())>1 else None
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, stratify=strat, random_state=42)

numeric_features = ['white_rating','black_rating','rating_diff','abs_rating_diff','moves','moves_log','material_diff','white_higher']
categorical_features = ['time_control','opening']

numeric_transformer = Pipeline([('scaler', StandardScaler())])
categorical_transformer = Pipeline([
    ('onehot', OneHotEncoder(handle_unknown='ignore', sparse_output=False))
])
preprocessor = ColumnTransformer([('num', numeric_transformer, numeric_features), ('cat', categorical_transformer, categorical_features)], remainder='drop')

dummy = Pipeline([('pre', preprocessor), ('clf', DummyClassifier(strategy='most_frequent'))])
dummy.fit(X_train, y_train)
y_dummy_pred = dummy.predict(X_test)
dummy_acc = accuracy_score(y_test, y_dummy_pred)

rf_pipeline = Pipeline([('pre', preprocessor), ('clf', RandomForestClassifier(n_estimators=150, random_state=42, class_weight='balanced'))])
rf_pipeline.fit(X_train, y_train)
y_pred_rf = rf_pipeline.predict(X_test)
rf_acc = accuracy_score(y_test, y_pred_rf)
rf_report = classification_report(y_test, y_pred_rf, output_dict=False)
cm_rf = confusion_matrix(y_test, y_pred_rf, labels=['white','draw','black'])

ohe = rf_pipeline.named_steps['pre'].named_transformers_['cat'].named_steps['onehot']
ohe_feature_names = list(ohe.get_feature_names_out(categorical_features))
feature_names = numeric_features + ohe_feature_names
importances = rf_pipeline.named_steps['clf'].feature_importances_
indices = np.argsort(importances)[::-1][:12]
top_feat = [feature_names[i] for i in indices]
top_imp = importances[indices]

plt.figure(figsize=(8,5)); plt.barh(range(len(top_imp))[::-1], top_imp[::-1]); plt.yticks(range(len(top_imp)), top_feat[::-1]); plt.title('Top features - Random Forest'); plt.tight_layout()
fig5 = os.path.join(FIG_DIR, "feat_imp_rf_pgn.png"); plt.savefig(fig5); plt.close()

use_xgb = True
try:
    import xgboost as xgb
except Exception:
    use_xgb = False

from sklearn.ensemble import GradientBoostingClassifier

xgb_pipeline = Pipeline([
    ('pre', preprocessor),
    ('clf', GradientBoostingClassifier(random_state=42))
])
xgb_pipeline.fit(X_train, y_train)
y_pred_xgb = xgb_pipeline.predict(X_test)

xgb_acc = accuracy_score(y_test, y_pred_xgb)
xgb_report = classification_report(y_test, y_pred_xgb, output_dict=False)
cm_xgb = confusion_matrix(y_test, y_pred_xgb, labels=['white','draw','black'])

joblib.dump(rf_pipeline, os.path.join(MODEL_DIR, "rf_pipeline_pgn.pkl"))
joblib.dump(xgb_pipeline, os.path.join(MODEL_DIR, "xgb_pipeline_pgn.pkl"))

plt.figure(figsize=(5,4)); plt.imshow(cm_rf, interpolation='nearest'); plt.title('Confusion Matrix - RF'); plt.colorbar()
ticks = np.arange(3); plt.xticks(ticks, ['white','draw','black']); plt.yticks(ticks, ['white','draw','black'])
for i in range(cm_rf.shape[0]):
    for j in range(cm_rf.shape[1]):
        plt.text(j, i, str(cm_rf[i,j]), ha='center', va='center')
plt.tight_layout(); cm_rf_file = os.path.join(FIG_DIR, "cm_rf_pgn.png"); plt.savefig(cm_rf_file); plt.close()

plt.figure(figsize=(5,4)); plt.imshow(cm_xgb, interpolation='nearest'); plt.title('Confusion Matrix - XGB/GB'); plt.colorbar()
ticks = np.arange(3); plt.xticks(ticks, ['white','draw','black']); plt.yticks(ticks, ['white','draw','black'])
for i in range(cm_xgb.shape[0]):
    for j in range(cm_xgb.shape[1]):
        plt.text(j, i, str(cm_xgb[i,j]), ha='center', va='center')
plt.tight_layout(); cm_xgb_file = os.path.join(FIG_DIR, "cm_xgb_pgn.png"); plt.savefig(cm_xgb_file); plt.close()

try:
    rf_cv = cross_val_score(rf_pipeline, X_train, y_train, cv=3, scoring='accuracy', n_jobs=-1)
    xgb_cv = cross_val_score(xgb_pipeline, X_train, y_train if not use_xgb else y_train_num, cv=3, scoring='accuracy', n_jobs=-1)
except Exception:
    rf_cv = None
    xgb_cv = None


# ---------- Código para gerar um relatório das partidas em .docx ----------
doc = Document()
doc.add_heading('PI1 - Predição do Resultado de Partidas de Xadrez (PGN local)', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph('Aluno: Matheus Franklin Brasileiro').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph('Data: ' + time.strftime("%d/%m/%Y")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_page_break()

doc.add_heading('1. Introdução', level=2)
doc.add_paragraph('Objetivo: prever o resultado de partidas de xadrez (white/draw/black) usando aprendizado de máquina supervisionado a partir de suas partidas exportadas em PGN.')

doc.add_heading('2. Dados e ETL', level=2)
doc.add_paragraph(f'Foram extraídas {len(df)} partidas do arquivo PGN: {PGN_FILE}.')
doc.add_paragraph('Foram criadas features: white_rating, black_rating, rating_diff, abs_rating_diff, moves, moves_log, opening, time_control, material_diff, white_higher.')

doc.add_heading('3. Análise Exploratória', level=2)
doc.add_paragraph('Distribuição dos resultados:'); doc.add_picture(fig1, width=Inches(5.5))
doc.add_paragraph('Rating diff por resultado:'); doc.add_picture(fig2, width=Inches(5.5))
doc.add_paragraph('Número de lances:'); doc.add_picture(fig3, width=Inches(5.5))

doc.add_heading('4. Modelos e Avaliação', level=2)
doc.add_paragraph(f'Baseline (Dummy) accuracy: {dummy_acc:.3f}')
doc.add_paragraph(f'Random Forest accuracy: {rf_acc:.3f}')
doc.add_paragraph('Relatório de classificação - Random Forest:'); doc.add_paragraph(rf_report)
doc.add_paragraph('Matriz de confusão - Random Forest:'); doc.add_picture(cm_rf_file, width=Inches(5.0))

doc.add_paragraph(f'XGBoost/GradientBoosting accuracy: {xgb_acc:.3f}')
doc.add_paragraph('Relatório de classificação - XGBoost/GB:'); doc.add_paragraph(xgb_report)
doc.add_paragraph('Matriz de confusão - XGBoost/GB:'); doc.add_picture(cm_xgb_file, width=Inches(5.0))

doc.add_heading('5. Importância das features', level=2)
doc.add_paragraph('Top features segundo Random Forest:'); doc.add_picture(fig5, width=Inches(5.5))

doc.add_heading('6. Conclusão', level=2)
doc.add_paragraph('Resumo: Modelos treinados com as partidas locais. Para robustez aumentar amostra e incluir features baseadas em sequência de jogadas (ex: número de blunders, tempo por jogada, motivos de vitória).')

doc.add_page_break()
doc.add_heading('Anexos', level=2)
doc.add_paragraph('CSV com dados extraídos: ' + csv_path)
doc.add_paragraph('Modelos salvos em: ' + MODEL_DIR)
doc.add_paragraph('Figuras em: ' + FIG_DIR)

out_doc = os.path.join(OUT_DIR, "PI1_Predicao_Xadrez_PGN_Matheus.docx")
doc.save(out_doc)
print("Relatório salvo em:", out_doc)
print("Arquivos gerados em:", OUT_DIR)
